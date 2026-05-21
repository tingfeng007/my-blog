#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
把 Word 版每日投资市场日报发布为网站在线阅读页。

用法:
    python generate_daily_report.py "D:/Downloads/每日投资市场日报 _ 2026-05-20.docx"
"""

from __future__ import annotations

import argparse
import html
import os
import re
import shutil
import sys
import zipfile
from pathlib import Path

try:
    from docx import Document
except ModuleNotFoundError:
    bundled_python = Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "python" / "python.exe"
    launched_as_script = Path(sys.argv[0]).resolve() == Path(__file__).resolve()
    if launched_as_script and bundled_python.exists() and Path(sys.executable).resolve() != bundled_python.resolve():
        os.execv(str(bundled_python), [str(bundled_python), str(Path(__file__).resolve()), *sys.argv[1:]])
    raise


ROOT = Path(__file__).resolve().parent
DOWNLOADS = ROOT / "downloads"
ASSETS = ROOT / "assets"


def infer_date(source: Path) -> str:
    match = re.search(r"(20\d{2})[-_.年 ]?(\d{1,2})[-_.月 ]?(\d{1,2})", source.stem)
    if not match:
        raise ValueError("文件名里没有找到日期，请使用类似 2026-05-20 的命名。")
    year, month, day = match.groups()
    return f"{year}-{int(month):02d}-{int(day):02d}"


def text_of(cell_or_paragraph) -> str:
    return html.escape((cell_or_paragraph.text or "").strip())


def extract_images(docx_path: Path, slug: str) -> list[str]:
    image_dir = ASSETS / slug
    image_dir.mkdir(parents=True, exist_ok=True)
    images: list[str] = []

    with zipfile.ZipFile(docx_path) as docx_zip:
        media = [name for name in docx_zip.namelist() if name.startswith("word/media/")]
        for index, name in enumerate(media, start=1):
            suffix = Path(name).suffix.lower() or ".png"
            target = image_dir / f"image-{index:02d}{suffix}"
            target.write_bytes(docx_zip.read(name))
            images.append(target.relative_to(ROOT).as_posix())

    return images


def table_html(table) -> str:
    rows = []
    for row_index, row in enumerate(table.rows):
        cells = []
        tag = "th" if row_index == 0 else "td"
        for cell in row.cells:
            cells.append(f"<{tag}>{text_of(cell)}</{tag}>")
        rows.append("<tr>" + "".join(cells) + "</tr>")
    return '<div class="report-table-wrap"><table class="report-table">' + "".join(rows) + "</table></div>"


def document_html(docx_path: Path, date: str, image_paths: list[str]) -> str:
    doc = Document(docx_path)
    blocks: list[str] = [f'<h1 class="report-title">每日投资市场日报 | {date}</h1>']
    image_index = 0

    for child in doc.element.body:
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            paragraph = next((p for p in doc.paragraphs if p._p is child), None)
            if paragraph is None:
                continue
            text = text_of(paragraph)
            if not text:
                continue
            style_name = (paragraph.style.name if paragraph.style and paragraph.style.name else "").lower()
            if "heading 1" in style_name or text.startswith(("一、", "二、", "三、", "四、", "五、", "六、")):
                blocks.append(f'<h2 class="report-h2">{text}</h2>')
            elif "heading 2" in style_name or re.match(r"^\d+[.、]", text):
                blocks.append(f'<h3 class="report-h3">{text}</h3>')
            elif text.startswith(("核心内容：", "核心逻辑：", "当日流向：", "风险提示：")):
                blocks.append(f'<p class="report-body emphasis-line">{text}</p>')
            elif text.startswith(("发布时间：", "来源：", "生成时间：")):
                blocks.append(f'<p class="report-meta-line">{text}</p>')
            else:
                blocks.append(f'<p class="report-body">{text}</p>')
        elif tag == "tbl":
            table = next((t for t in doc.tables if t._tbl is child), None)
            if table is not None:
                blocks.append(table_html(table))
                if image_index < len(image_paths):
                    src = image_paths[image_index]
                    blocks.append(f'<figure class="report-figure"><img src="{src}" alt="日报图表"><figcaption>文档图表</figcaption></figure>')
                    image_index += 1

    for src in image_paths[image_index:]:
        blocks.append(f'<figure class="report-figure"><img src="{src}" alt="日报图表"><figcaption>文档图表</figcaption></figure>')

    return "\n".join(blocks)


def page_template(date: str, slug: str, docx_name: str, content: str) -> str:
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <meta name="description" content="每日投资市场日报，{date}，在线阅读版。">
  <meta property="og:title" content="每日投资市场日报 | {date}">
  <meta property="og:description" content="听风整理的每日投资市场观察，保留 Word 原文结构、表格和图示。">
  <meta property="og:type" content="article">
  <meta property="og:image" content="assets/site/market-cover.jpg">
  <meta name="twitter:card" content="summary_large_image">
  <title>每日投资市场日报 | {date}</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@500;600;700;900&family=Noto+Sans+SC:wght@300;400;500;700&display=swap" rel="stylesheet">
  <link rel="icon" href="data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'%3E%3Crect width='64' height='64' fill='%23191813'/%3E%3Ctext x='32' y='43' font-size='34' text-anchor='middle' fill='%23fbfaf6' font-family='serif'%3E%E5%90%AC%3C/text%3E%3C/svg%3E">
  <link rel="stylesheet" href="style.css">
</head>
<body class="report-page">
  <nav class="nav" aria-label="主导航">
    <div class="nav-container">
      <a href="index.html" class="logo">文明切片</a>
      <button class="nav-toggle" type="button" aria-label="打开导航" aria-expanded="false" aria-controls="primary-navigation"><em>菜单</em><span></span><span></span></button>
      <div class="nav-links" id="primary-navigation">
        <a href="index.html#essays" class="nav-link">文章</a>
        <a href="market.html" class="nav-link">市场日报</a>
        <a href="index.html#signals" class="nav-link">资讯切片</a>
        <a href="ai-hub.html" class="nav-link">AI 启动台</a>
        <a href="about.html" class="nav-link">关于</a>
      </div>
    </div>
  </nav>
  <main class="report-shell">
    <div class="report-toolbar">
      <a href="market.html">返回日报归档</a>
      <span>在线原文阅读</span>
      <button type="button" onclick="shareContent('每日投资市场日报 - {date}', '{slug}.html')">分享</button>
      <a href="downloads/{docx_name}">下载 Word 原文</a>
    </div>
    <article class="report-document">
{content}
    </article>
  </main>
  <footer class="footer"><p class="footer-text">© 2026 文明切片 · 听风. All Rights Reserved.</p></footer>
  <script src="script.js"></script>
</body>
</html>
"""


def first_summary(content: str) -> str:
    match = re.search(r'<p class="report-body(?: emphasis-line)?">(.+?)</p>', content)
    text = re.sub(r"<[^>]+>", "", html.unescape(match.group(1))) if match else "听风整理的每日投资市场观察。"
    return text[:86] + ("…" if len(text) > 86 else "")


def replace_between(text: str, start: str, end: str, replacement: str) -> str:
    pattern = re.compile(re.escape(start) + r".*?" + re.escape(end), re.S)
    return pattern.sub(start + "\n" + replacement.strip() + "\n" + end, text)


def update_indexes(date: str, slug: str, summary: str, docx_name: str) -> None:
    display_date = date.replace("-", ".")
    home_card = f"""          <article class="market-card reveal-on-scroll">
            <div class="market-date">{display_date} <span>最新</span></div>
            <h3>每日投资市场日报</h3>
            <p>{html.escape(summary)}</p>
            <div class="card-meta">
              <button type="button" onclick="shareContent('每日投资市场日报 - {display_date}', '{slug}.html')">分享</button>
              <a href="{slug}.html">在线阅读</a>
            </div>
          </article>"""
    archive_item = f"""        <article class="archive-item">
          <time>{display_date}</time>
          <div>
            <h2>每日投资市场日报</h2>
            <p>{html.escape(summary)}</p>
          </div>
          <div class="archive-actions">
            <a href="{slug}.html">在线阅读</a>
            <a href="downloads/{docx_name}">下载原文</a>
          </div>
        </article>"""

    index_path = ROOT / "index.html"
    index_text = index_path.read_text(encoding="utf-8")
    if f'href="{slug}.html"' not in index_text:
        block = re.search(r"<!-- MARKET_CARDS_START -->(.*?)<!-- MARKET_CARDS_END -->", index_text, re.S).group(1)
        cards = re.findall(r"\s*<article class=\"market-card.*?</article>", block, re.S)
        new_cards = "\n".join([home_card] + cards[:2])
        index_path.write_text(replace_between(index_text, "<!-- MARKET_CARDS_START -->", "<!-- MARKET_CARDS_END -->", new_cards), encoding="utf-8")

    archive_path = ROOT / "market.html"
    archive_text = archive_path.read_text(encoding="utf-8")
    if f'href="{slug}.html"' not in archive_text:
        archive_text = archive_text.replace("<!-- MARKET_ARCHIVE_START -->", "<!-- MARKET_ARCHIVE_START -->\n" + archive_item)
        archive_path.write_text(archive_text, encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(description="把 Word 日报转换为网站在线阅读页，并更新首页和日报归档。")
    parser.add_argument("docx", help="Word 日报路径，例如 D:/Downloads/每日投资市场日报 _ 2026-05-20.docx")
    args = parser.parse_args()

    source = Path(args.docx).expanduser().resolve()
    if not source.exists() or source.suffix.lower() != ".docx":
        raise FileNotFoundError(f"找不到 Word 文档: {source}")

    date = infer_date(source)
    slug = f"market-daily-{date}"
    docx_name = f"{slug}.docx"

    DOWNLOADS.mkdir(exist_ok=True)
    shutil.copy2(source, DOWNLOADS / docx_name)

    image_paths = extract_images(source, slug)
    content = document_html(source, date, image_paths)
    (ROOT / f"{slug}.html").write_text(page_template(date, slug, docx_name, content), encoding="utf-8")
    update_indexes(date, slug, first_summary(content), docx_name)

    print(f"已生成: {slug}.html")
    print(f"已复制: downloads/{docx_name}")
    print("已更新: index.html, market.html")


if __name__ == "__main__":
    main()
